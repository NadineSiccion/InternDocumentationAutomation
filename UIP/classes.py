# program that generates Daily Journals for UIP
import os
import json
from datetime import date, timedelta
from docx import Document
from docx.shared import Inches
from pathlib import Path




class Internship():
    def __init__(self, name, remaining_hours:int, djnum_counter):
        self.name = name
        self.remaining_hours = remaining_hours
        self.djnum_counter = djnum_counter
    
    def save_internship(self):
        data_load = {
            'name': self.name,
            'remaining_hours': self.remaining_hours,
            'djnum_counter': self.djnum_counter
            }
        with open(f'{self.name}_internship_save.json', 'w') as save_file:
            json.dump(data_load, save_file, indent=4)
        print(f'A save file for {self.name} has been written.')        
        
    def get_daily_entry(self):
        title = ''
        journal_entry = ''

        with open (f'{self.name}_input.txt', 'r') as in_file:
            while True:
                line = in_file.readline()
                if 'Title' in line:
                    print('title to perform')
                    split_line = line.split(':')[1:]
                    for item in split_line:
                        item = item.strip()
                    title = ' '.join(split_line).strip()
                    print('title done')
                elif 'Description' in line:
                    print('desc to perform')
                    desc_line = 'a'
                    while not (desc_line.strip() == ''): 
                        desc_line = in_file.readline()
                        print(desc_line)
                        if desc_line:
                            if desc_line[-1] == '\n':
                                journal_entry += '\n'
                        journal_entry += (desc_line.strip() + '\n')
                        print('journal entry: ' + journal_entry)
                    print('desc done')
                    break
        return {'title': title, 'entry': journal_entry}


        # title
        # journal_desc
        # with open (input_daily_path, 'r') as intern_save:
        #     data = intern_save.readlines()
        #     for line in data:
        #         if 'Title' in line:
        #             title_line = line.split(':')
        #             title = title_line[1:].strip()
        #         elif 'Description' in line:
        #             desc_line = line.split(':')
        #             journal_desc = desc_line[1:]
    
    def update_djnum(self, djupdate=1):
        self.djnum_counter += djupdate
        
    def update_hours(self):
        self.remaining_hours -= 8
    
    def penalize(self, penalty:int):
        self.remaining_hours += penalty
    
    def set_hours_to(self, hours):
        self.remaining_hours = hours

    def load_internship(name):
        try:
            with open(f'{name}_internship_save.json', 'r') as save_file:
                data = json.load(save_file)
        except Exception as e:
            print(f"An error occurred: {e}")
            return 0
        
        remaining_hours = data['remaining_hours']
        djnum = data['djnum_counter']
        internship = Internship(name, remaining_hours, djnum)

        return internship

class Report(Internship):
    def __init__(self, name, remaining_hours:int, djnum_counter, template_path:Path, output_path:Path, image_path:Path, report_date:date):
        super().__init__(name, remaining_hours, djnum_counter)
        self.template_path = template_path
        self.output_path = output_path
        self.image_path = image_path
        self.report_date = report_date
        self.cwd = os.getcwd()

    def set_date(self, date):
        self.date = date
    
    def generate_data(self)->dict:
        td = self.get_daily_entry()

        djnum = self.djnum_counter
        month = self.report_date.strftime('%B')
        day = self.report_date.day
        year = self.report_date.year
        remaining_hours = self.remaining_hours
        journal_title = td['title']
        journal_desc = td['entry']

        data_dict = {
        '[DJNUM]': str(djnum),
        '[MONTH]': str(month),
        '[DAY]': str(day),
        '[YEAR]': str(year),
        '[REMAINING HOURS]': str(remaining_hours),
        '[JOURNAL TITLE]': journal_title,
        '[JOURNAL DESCRIPTION]': journal_desc
        }
        return data_dict
    
    def add_to_current_day(self, days=1):
        if self.report_date:
            self.report_date += timedelta(days=days)
        else:
            print("No date assigned to this Report object.")
    
    def save_report(self):
        # Unsure that paths are relative and .relative_to is not used in a redundant manner
        if self.template_path.is_absolute():
            save_template_path = str(self.template_path.relative_to(self.cwd))
        else:
            save_template_path = str(self.template_path)
        
        if self.output_path.is_absolute():
            save_output_path = str(self.output_path.relative_to(self.cwd))
        else:
            save_output_path = str(self.output_path)
        
        if self.image_path.is_absolute():
            save_image_path = str(self.image_path.relative_to(self.cwd))
        else:
            save_image_path = str(self.image_path)

        report_dict = {
        'name': self.name,
        'remaining_hours': self.remaining_hours,
        'djnum_counter': self.djnum_counter,
        'template_path': save_template_path,
        'output_path': save_output_path,
        'image_path': save_image_path,
        'report_date': {
            'year': str(self.report_date.year), 
            'month': str(self.report_date.month),
            'day': str(self.report_date.day)}
        }
        file_name = f'{self.name}_report.json'

        with open(file_name, 'w') as save_file:
            json.dump(report_dict, save_file, indent=4)
        print(f'Save file for {self.name} Report has been created')
    
    def add_absent(self, consecutive_absents = 1):
        self.djnum_counter -= consecutive_absents
        print('Deducted ' + consecutive_absents + ' from djnum')
 
    def edit_document(self):
        doc = Document(self.template_path)
        data = self.generate_data()

        img_file1 = self.image_path.joinpath(f'd{str(data["[DJNUM]"]) + ' (1).PNG'}')
        img_file2 = self.image_path.joinpath(f'd{str(data["[DJNUM]"]) + ' (2).PNG'}')

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(key, value)
                        print('Replaced "' + key + '" with "' + value + '".')
        
        print(f'looking for file "{str(img_file1)}"')
        doc.add_picture(str(img_file1), width=Inches(6))
        print('Picture added.')

        print(f'looking for file "{str(img_file2)}"')
        doc.add_picture(str(img_file2), width=Inches(6))
        print('Picture added.')
        
        new_filename = str(self.output_path.name).replace("###", f"{self.djnum_counter:03d}")
        doc.save(self.output_path.with_name(new_filename))
        print('Saved output at ' + str(self.output_path).replace('###', f'{self.djnum_counter:03d}') + '.')
    
    def load_report(name):
        """Loads report data from a save file"""
        try:
            with open(f'{name}_report.json', 'r') as save_file:
                data = json.load(save_file)
        except Exception as e:
            print(f"An error occurred: {e}")
            return 0
        
        remaining_hours = data['remaining_hours']
        djnum = data['djnum_counter']
        template_path = Path(data['template_path'])
        output_path = Path(data['output_path'])
        image_path = Path(data['image_path'])
        report_date = date(int(data['report_date']['year']), int(data['report_date']['month']), int(data['report_date']['day']))
        
        report = Report(name, remaining_hours, djnum, template_path, output_path, image_path, report_date)
        

        return report
    
    def update_counters(self, day:str):
        '''Updates djnum, hours and current date. To be done after generating the last report.
        Parameters: day -  can be "T" or "Th". Put the day of the report recently made.
        '''
        self.update_djnum(1)
        self.update_hours()
        if day.lower() == "t":
            self.add_to_current_day(2)
        elif day.lower() == "th":
            self.add_to_current_day(5)
        
        print('Report counters updated for the next report.')



# edit_document(template_path, output_path, image_path, data)

# #variables
# djnum = str(2)
# month = 'September'
# day = '2'
# year = '2024'
# remaining_hours = '345'
# journal_title = 'What I learned today'
# journal_desc = '''wow, I really learned a lot.

# text text and more text.'''

# data = {
#     '[DJNUM]': djnum,
#     '[MONTH]': month,
#     '[DAY]': day,
#     '[YEAR]': year,
#     '[REMAINING HOURS]': remaining_hours,
#     '[JOURNAL TITLE]': journal_title,
#     '[JOURNAL DESCRIPTION]': journal_desc
# }