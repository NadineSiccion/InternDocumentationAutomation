from classes import Report
from docx import Document
from docx.shared import Inches
from pathlib import Path
from datetime import date


def edit_document(template_path, output_path, image_path, data:dict):
    doc = Document(template_path)

    img_file1 = image_path + f'd{data["[DJNUM]"] + ' (1).PNG'}'
    img_file2 = image_path + f'd{data["[DJNUM]"] + ' (2).PNG'}'

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, value)
                    print('Replaced "' + key + '" with "' + value + '".')
    
    print(f'looking for file "{img_file1}"')
    doc.add_picture(img_file1, width=Inches(6))
    print('Picture added.')

    print(f'looking for file "{img_file2}"')
    doc.add_picture(img_file2, width=Inches(6))
    print('Picture added.')
    
    doc.save(output_path)
    print('Saved output at ' + str(output_path) + '.')


# run if main
if __name__ == '__main__':
    # days from Tue to Thur = 2
    # days from Thur to Tue = 5

    cwd = Path.cwd()


    template_path = cwd /'templates'/'DAILY_JOURNAL_TEMPLATE.docx'
    output_path = cwd /'outputs'/'DailyJounalNo###.docx'
    image_path = cwd / 'imgs'
    report_date = date(2024, 9, 26)

    # CODE FIRST SCRIPT RUN FOR DJ#6 --DO NOT RUN
    # image file names must follow the format "d# (1).PNG"
    # uip_internship_report = Report('uip', 561, 6, template_path, output_path, image_path, report_date)

    # uip_internship_report.update_hours()
    # uip_internship_report.penalize(1)
    # uip_internship_report.save_report()
    # uip_internship_report.edit_document()

    # DAY 7
    # uip_internship_report = Report.load_report('uip')
    # # print(type(uip_internship_report), isinstance(uip_internship_report, Report))
    # uip_internship_report.update_counters('Th')
    # uip_internship_report.penalize(1)
    # uip_internship_report.save_report()
    # uip_internship_report.edit_document()

    # use .update() with "T" to add 2 days, and "Th" to add 5 days

    uip_internship_report = Report.load_report('uip')
    uip_internship_report.update_counters('T')
    uip_internship_report.edit_document()
    uip_internship_report.save_report()
