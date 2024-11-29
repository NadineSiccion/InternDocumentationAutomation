import os
import json
from datetime import date, timedelta
from docx import Document
from docx.shared import Inches
from pathlib import Path
from docx.shared import Pt
import inspect

# FIXME: make so date is really only a string during saving and loading, any other time, it should be a date

class Internship():
    def __init__(self, week_no:int, required_hrs:int, completed_hrs:int, start_date:date, current_week_start):
        self.week_no = int(week_no)
        self.required_hrs = int(required_hrs)
        self.completed_hrs = int(completed_hrs)
        self.start_date = Internship.str_to_date(start_date)
        self.current_week_start = Internship.str_to_date(current_week_start)

    def date_to_str(date:date) -> str:
        return date.strftime("%B %d, %Y")
    
    def str_to_date(datestring:str) -> date:
        temp = str(datestring).split('-')
        return date(int(temp[0]), int(temp[1]), int(temp[2]))
    
    def get_week_dates(self, week_start:date)->list[date]:
        """Returns a list of 5 consecutive days starting from a given date YYYY/MM/DD"""
        temp_list = []
        for i in range(5):
            temp_list.append(week_start)
            week_start += timedelta(days=1)
        
        week_dates = []
        for day in temp_list:
            week_dates.append(Internship.date_to_str(day))
        return week_dates
    
    def get_remaining_hrs(self) -> int:
        """Return Required_Hours - Completed_Hours"""
        return self.required_hrs - self.completed_hrs
    
    def add_absent(self, hours:int = 6) -> None:
        self.completed_hrs -= hours
        print('Deducted ' + str(hours) + 
              ' hours from completed hours.')
    
    def add_hour_difference(self, hours:int) -> None:
        self.completed_hrs += hours
        if hours > 0:
            print("Added " + str(hours) + 
                  " hours to completed hours")
        else:
            print('Deducted ' + str(hours) + 
              ' hours from completed hours.')
    
    def update_week_no(self, num:int = 1) -> None:
        self.week_no += num
        print('Week no. incremented by ' + str(num))

    def compute_week_number(self, week_start):
        week_start = Internship.str_to_date(week_start)
        week_no = 0
        counter_week = self.start_date
        while counter_week < week_start:
            counter_week += timedelta(weeks=1)
            week_no += 1
        
        return week_no
    
    def update_internship(self):
        # add 1 to week_no
        self.week_no += 1
        # add 30 hrs to completed hrs
        self.completed_hrs += (6*5)
        # add 1 week to current start week
        self.current_week_start += timedelta(weeks=1)
        print('Updated data from Week No. ',self.week_no -1, ' to ', self.week_no, 
              '. Ready to generate next report.')

class Report(Internship):
    def __init__(self, week_no:int, required_hrs:int, completed_hrs:int, start_date:date, current_week_start, template_file:Path, out_file:Path, save_file:Path):
        super().__init__(week_no, required_hrs, completed_hrs, start_date, current_week_start)
        self.template_file = template_file
        self.out_file = out_file
        self.save_file = save_file
        
    # == SAVE AND LOAD FUNCTIONS ==

    def save_report(self) -> None:
        """Saves the current report object as a json in save location."""
        save_data = {
            'week_no': str(self.week_no), 
            'required_hrs': str(self.required_hrs), 
            'completed_hrs': str(self.completed_hrs), 
            'start_date': str(self.start_date),
            'current_week_start': str(self.current_week_start),
            'template_file': str(self.template_file),
            'out_file': str(self.out_file),
            'save_file': str(self.save_file)
        }
        
        print("Generating save data...")
        for key, value in save_data.items():
            print(key, ": ", type(value))
        print('Save data generated.')
        
        with open(self.save_file, "w") as file:
            json.dump(save_data, file, indent=4)
        print('Report data saved to ', self.save_file, ".")
    
    def load_report(file_name): #FIXME
        """Loads report data from a save file. Returns Report object."""
        try:
            with open(file_name, 'r') as save_file:
                data = json.load(save_file)
        except Exception as e:
            print(f"An error occurred: {e}")
            return 0
        
        
        week_no = int(data['week_no']) 
        required_hrs = int(data['required_hrs']) 
        completed_hrs = int(data['completed_hrs']) 
        start_date = Internship.str_to_date(data['start_date'])
        current_week_start = Internship.str_to_date(data['current_week_start'])
        template_file = Path(data['template_file'])
        out_file = Path(data['out_file'])
        save_file = Path(data['save_file'])
        
        report = Report(
            week_no, 
            required_hrs, 
            completed_hrs, 
            start_date, 
            current_week_start, 
            template_file, 
            out_file, 
            save_file)
        
        return report
    
    # == REPORT GENERATING FUNCTIONS ==

    def generate_data(self, week_start:str = None, hours_per_day:int = 6, days_per_week:int = 5) -> dict:
        """Returns a dict of the data to be used in making the new document"""
        self.current_week_start = week_start

        start_date = self.start_date
        if (week_start != None):
            week_start = Internship.str_to_date(week_start)
        else: 
            week_start = self.current_week_start
        weekdays = self.get_week_dates(week_start)
        week_end = weekdays[-1]

        week_no = self.week_no
        week_hrs = hours_per_day*days_per_week
        week_mins = week_hrs*60
        
        data_dict = {
        '{week_no}': str(week_no),
        '{week_start}': Internship.date_to_str(week_start),
        '{week_end}': str(week_end),
        '{day_1}': weekdays[0],
        '{day_2}': weekdays[1],
        '{day_3}': weekdays[2],
        '{day_4}': weekdays[3],
        '{day_5}': weekdays[4],
        '{week_hrs}': str(week_hrs),
        '{week_mins}': f'{week_mins:,}',
        '{remaining_hrs}': str(self.get_remaining_hrs())
        }

        print('Output: \n')
        for key, value in data_dict.items():
            print(f"{key}: {value}")

        return data_dict

    def turn_bold(paragraph):
        for run in paragraph.runs:
            run.bold = True
    
    def edit_document(self, days_in_week=None):
        doc = Document(self.template_file)
        if days_in_week == None: 
            data = self.generate_data(str(self.current_week_start))
        else:
            data = self.generate_data(str(self.current_week_start), days_per_week=days_in_week)

        style = doc.styles['Normal']
        font = style.font
        font.name = "Calibri Light"

        lg_style = doc.styles.add_style('Larger', 1)
        lg_font = lg_style.font
        lg_font.name = "Calibri Light"
        lg_font.size = Pt(14)

        for row in doc.tables:
            for cell in row._cells:
                for key, value in data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            print('Replaced "' + key + '" with "' + value + '" in table.')
                            for paragraph in cell.paragraphs:
                                paragraph.style = doc.styles['Normal']
                                if key in ['{week_no}', '{week_end}','{remaining_hrs}']:
                                    paragraph.style = doc.styles['Larger']
                                Report.turn_bold(paragraph)
        
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(key, value)
                            print('Replaced "' + key + '" with "' + value + '" in paragraph.')
                        paragraph.style = doc.styles['Normal']
                        Report.turn_bold(paragraph)
                    
            # for key, value in data.items():
            #     if key in paragraph.text:
            #             for run in paragraph.runs:
            #                 run.text = run.text.replace(key, value)
            #             print('Replaced "' + key + '" with "' + value + '".')
        
        new_filename = str(self.out_file.name).replace("###", f"{self.week_no:03d}")
        with open(new_filename, 'w') as out_file:
            doc.save(self.out_file.with_name(new_filename)) # There's a bug here but the sofware works fine
        
        print('Saved output at ' + str(new_filename) + '.')
